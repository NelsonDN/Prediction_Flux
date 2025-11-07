"""
GENERATEUR DE RAPPORT DE PRESENTATION PROFESSIONNELLE
Rapport sur l'optimisation des flux de tresorerie bancaire
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill_color):
    """Definir la couleur de fond d'une cellule"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_red_heading(doc, text, level=1):
    """Ajouter un titre en rouge"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(139, 0, 0)  # Rouge fonce
    return heading


def generate_presentation_report():
    """Genere le rapport de presentation professionnelle"""
    
    doc = Document()
    
    # Configuration des styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ========================================================================
    # PAGE 1 : PAGE DE TITRE
    # ========================================================================
    
    # Titre principal
    title = doc.add_heading('OPTIMISATION DES FLUX DE TRESORERIE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(139, 0, 0)
        run.font.size = Pt(24)
    
    # Sous-titre
    subtitle = doc.add_paragraph('Systeme d\'Allocation Optimale par Programmation Lineaire')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(139, 0, 0)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Informations document
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Rapport Technique\n').bold = True
    info.add_run('Departement Tresorerie\n')
    
    doc.add_page_break()
    
    # ========================================================================
    # PAGE 2 : APPROCHE GLOBALE
    # ========================================================================
    
    add_red_heading(doc, '1. APPROCHE GLOBALE DU PROJET', level=1)
    
    # Objectif
    doc.add_heading('1.1 Objectif', level=2)
    doc.add_paragraph(
        'Optimiser les flux de tresorerie entre 29 agences bancaires et la BEAC '
        'en minimisant les couts de transport tout en garantissant l\'equilibre '
        'entre besoins et excedents de liquidite.'
    )
    
    # Methodologie
    doc.add_heading('1.2 Methodologie en deux phases', level=2)
    
    doc.add_paragraph()
    p = doc.add_paragraph('Phase 1 : Prevision des besoins', style='List Bullet')
    p.runs[0].font.bold = True
    
    doc.add_paragraph(
        'Prediction des besoins en liquidite sur 5 jours (J+1 a J+5) par '
        'Machine Learning avec evaluation de 6 modeles et selection automatique '
        'du meilleur modele selon le critere MAPE minimal.'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph('Phase 2 : Optimisation des flux', style='List Bullet')
    p.runs[0].font.bold = True
    
    doc.add_paragraph(
        'Allocation optimale des flux monetaires entre agences par programmation '
        'lineaire, avec minimisation des couts de transport et generation automatique '
        'de plans d\'action quotidiens.'
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # PAGE 3 : METHODE D'OPTIMISATION CBC/SIMPLEXE
    # ========================================================================
    
    add_red_heading(doc, '2. METHODE D\'OPTIMISATION : CBC ET SIMPLEXE', level=1)
    
    # Nature du probleme
    doc.add_heading('2.1 Type de probleme', level=2)
    doc.add_paragraph(
        'Le systeme resout un probleme de transport classique en recherche operationnelle. '
        'Il s\'agit d\'allouer des ressources (liquidites) depuis des sources (agences '
        'en excedent) vers des destinations (agences en besoin) en minimisant le cout total.'
    )
    
    # Solveur CBC
    doc.add_heading('2.2 Solveur CBC (COIN-OR Branch and Cut)', level=2)
    
    doc.add_paragraph(
        'CBC est un solveur open-source de programmation lineaire developpe par le projet '
        'COIN-OR (Computational Infrastructure for Operations Research). Il combine deux '
        'techniques puissantes :'
    )
    
    doc.add_paragraph()
    
    # Simplexe
    p = doc.add_paragraph('A. Algorithme du Simplexe', style='List Number')
    p.runs[0].font.bold = True
    
    doc.add_paragraph(
        'Le simplexe explore les sommets du polyedre des solutions realisables. '
        'Il se deplace de sommet en sommet en ameliorant systematiquement la fonction '
        'objectif jusqu\'a atteindre l\'optimum.'
    )
    
    doc.add_paragraph('Principe de fonctionnement :')
    
    doc.add_paragraph('1. Partir d\'une solution realisable initiale', style='List Number 2')
    doc.add_paragraph('2. Identifier une direction d\'amelioration', style='List Number 2')
    doc.add_paragraph('3. Se deplacer le long de cette direction', style='List Number 2')
    doc.add_paragraph('4. Repeter jusqu\'a ce qu\'aucune amelioration ne soit possible', style='List Number 2')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Exemple visuel : Imaginez un diamant (polyedre). Le simplexe commence a un coin, '
        'se deplace d\'arrete en arrete vers les coins adjacents ayant une valeur superieure, '
        'jusqu\'a atteindre le sommet (solution optimale).'
    )
    
    doc.add_paragraph()
    
    # Branch-and-Bound
    p = doc.add_paragraph('B. Technique Branch-and-Bound', style='List Number')
    p.runs[0].font.bold = True
    
    doc.add_paragraph(
        'Cette technique divise le probleme en sous-problemes (branching) et elimine '
        'intelligemment les branches non prometteuses (bounding).'
    )
    
    doc.add_paragraph('Fonctionnement :')
    
    doc.add_paragraph(
        '1. Resolution du probleme relaxe (sans contraintes d\'integralite)',
        style='List Number 2'
    )
    doc.add_paragraph(
        '2. Si solution fractionnaire : diviser en deux sous-problemes',
        style='List Number 2'
    )
    doc.add_paragraph(
        '3. Eliminer les branches dont la borne est pire que la meilleure solution',
        style='List Number 2'
    )
    doc.add_paragraph(
        '4. Continuer jusqu\'a obtenir une solution entiere optimale',
        style='List Number 2'
    )
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Exemple visuel : Comme un arbre de decision ou chaque branche represente un choix. '
        'L\'algorithme coupe les branches qui ne peuvent mener a une meilleure solution.'
    )
    
    # Comparaison
    doc.add_heading('2.3 Comparaison avec autres methodes', level=2)
    
    # Tableau comparatif
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Headers
    headers = ['Methode', 'Optimalite', 'Complexite', 'Utilisation']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_background(cell, 'DC143C')  # Rouge
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Donnees
    data = [
        ['Nord-Ouest', 'Non (Heuristique)', 'O(n) - Rapide', 'Solution initiale'],
        ['Stepping-Stone', 'Oui', 'O(n²) - Lent', 'Amelioration iterative'],
        ['Balas-Hammer', 'Oui', 'Exponentiel', 'Petits problemes'],
        ['CBC/Simplexe', 'Oui (Garanti)', 'O(n³) - Rapide', 'Production industrielle']
    ]
    
    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value
            if i == 4:  # Derniere ligne (notre methode)
                table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Avantages CBC
    doc.add_heading('2.4 Avantages de CBC', level=2)
    
    doc.add_paragraph('Optimalite garantie : Solution mathematiquement prouvee comme minimale', style='List Bullet')
    doc.add_paragraph('Rapidite : Resolution en moins de 1 seconde pour 30 agences', style='List Bullet')
    doc.add_paragraph('Robustesse : Gestion automatique des contraintes complexes', style='List Bullet')
    doc.add_paragraph('Standard industriel : Utilise en aviation, logistique, finance', style='List Bullet')
    doc.add_paragraph('Open-source : Aucun cout de licence', style='List Bullet')
    
    doc.add_page_break()
    
    # ========================================================================
    # PAGE 4 : FORMULATION MATHEMATIQUE
    # ========================================================================
    
    add_red_heading(doc, '3. FORMULATION MATHEMATIQUE DU PROBLEME', level=1)
    
    # Variables
    doc.add_heading('3.1 Variables de decision', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('x[i,j]')
    run.font.italic = True
    p.add_run(' = Montant transfere de l\'agence i (excedent) vers l\'agence j (besoin)')
    
    doc.add_paragraph()
    
    # Parametres
    doc.add_heading('3.2 Parametres d\'entree', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('Excedent[i]')
    run.font.italic = True
    p.add_run(' : Montant disponible a l\'agence i')
    
    p = doc.add_paragraph()
    run = p.add_run('Besoin[j]')
    run.font.italic = True
    p.add_run(' : Montant requis par l\'agence j')
    
    p = doc.add_paragraph()
    run = p.add_run('C[i,j]')
    run.font.italic = True
    p.add_run(' : Cout unitaire pour transferer 1 million de i vers j (30-100 FCFA)')
    
    doc.add_paragraph()
    
    # Fonction objectif
    doc.add_heading('3.3 Fonction objectif', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('Minimiser Z = ')
    run.font.bold = True
    run = p.add_run('Somme(i) Somme(j) C[i,j] × x[i,j]')
    run.font.italic = True
    
    doc.add_paragraph()
    doc.add_paragraph('Traduction : Minimiser le cout total de tous les transferts inter-agences.')
    
    doc.add_paragraph()
    
    # Contraintes
    doc.add_heading('3.4 Contraintes', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Contrainte 1 : Conservation des excedents').bold = True
    
    p = doc.add_paragraph()
    run = p.add_run('Pour tout i : Somme(j) x[i,j] = Excedent[i]')
    run.font.italic = True
    
    doc.add_paragraph('Chaque agence en excedent doit transferer exactement tout son excedent.')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Contrainte 2 : Satisfaction des besoins').bold = True
    
    p = doc.add_paragraph()
    run = p.add_run('Pour tout j : Somme(i) x[i,j] = Besoin[j]')
    run.font.italic = True
    
    doc.add_paragraph('Chaque agence en besoin doit recevoir exactement son besoin.')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Contrainte 3 : Non-negativite').bold = True
    
    p = doc.add_paragraph()
    run = p.add_run('Pour tout i,j : x[i,j] >= 0')
    run.font.italic = True
    
    doc.add_paragraph('Les montants transferes doivent etre positifs ou nuls.')
    
    doc.add_paragraph()
    
    # Exemple
    doc.add_heading('3.5 Exemple illustratif', level=2)
    
    doc.add_paragraph('Situation :')
    doc.add_paragraph('Agence A : Excedent 100 millions', style='List Bullet')
    doc.add_paragraph('Agence B : Excedent 50 millions', style='List Bullet')
    doc.add_paragraph('Agence X : Besoin 80 millions', style='List Bullet')
    doc.add_paragraph('Agence Y : Besoin 70 millions', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Matrice des couts unitaires (FCFA par million) :')
    
    # Tableau exemple
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid'
    
    table.rows[0].cells[0].text = ''
    table.rows[0].cells[1].text = 'X'
    table.rows[0].cells[2].text = 'Y'
    
    table.rows[1].cells[0].text = 'A'
    table.rows[1].cells[1].text = '60'
    table.rows[1].cells[2].text = '80'
    
    table.rows[2].cells[0].text = 'B'
    table.rows[2].cells[1].text = '50'
    table.rows[2].cells[2].text = '70'
    
    doc.add_paragraph()
    doc.add_paragraph('Solution optimale trouvee par CBC :')
    
    doc.add_paragraph('A verse 30M a X (cout = 60 × 30 = 1 800 FCFA)', style='List Bullet')
    doc.add_paragraph('A verse 70M a Y (cout = 80 × 70 = 5 600 FCFA)', style='List Bullet')
    doc.add_paragraph('B verse 50M a X (cout = 50 × 50 = 2 500 FCFA)', style='List Bullet')
    doc.add_paragraph('B verse 0M a Y (pas de transfert)', style='List Bullet')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Cout total optimal : 9 900 FCFA').bold = True
    
    doc.add_page_break()
    
    # ========================================================================
    # PAGE 5 : ARCHITECTURE ET MISE EN PRODUCTION
    # ========================================================================
    
    add_red_heading(doc, '4. ARCHITECTURE DU SYSTEME', level=1)
    
    doc.add_heading('4.1 Etapes du pipeline', level=2)
    
    # Tableau pipeline
    table = doc.add_table(rows=10, cols=2)
    table.style = 'Light Grid'
    
    # Headers
    headers_row = table.rows[0]
    set_cell_background(headers_row.cells[0], 'DC143C')
    set_cell_background(headers_row.cells[1], 'DC143C')
    headers_row.cells[0].text = 'Etape'
    headers_row.cells[1].text = 'Description'
    for cell in headers_row.cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Donnees
    pipeline = [
        ('1. Selection modele', 'Evaluation 6 modeles ML, selection automatique du meilleur'),
        ('2. Matrice couts', 'Generation matrice symetrique 30×30 (30-100 FCFA/M)'),
        ('3. Donnees entree', 'Calcul flux optimal par agence (Besoin + Stock - Solde)'),
        ('4. Separation', 'Classification Besoin/Excedent, verification equilibre'),
        ('5. Optimisation CBC', 'Resolution probleme transport par simplexe'),
        ('6. Calcul couts', 'Application couts unitaires sur allocation optimale'),
        ('7. Export Excel', '5 fichiers (J+1 a J+5) avec 4 feuilles chacun'),
        ('8. Rapport Word', 'Synthese consolidee avec formulation mathematique'),
        ('9. Validation', 'Verification coherence et generation statistiques')
    ]
    
    for i, (etape, desc) in enumerate(pipeline, 1):
        table.rows[i].cells[0].text = etape
        table.rows[i].cells[1].text = desc
    
    doc.add_paragraph()
    
    # Logiciels requis
    doc.add_heading('4.2 Environnement technique', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Langage de programmation :').bold = True
    p.add_run(' Python 3.8 ou superieur')
    
    doc.add_paragraph()
    doc.add_paragraph('Bibliotheques Python requises :')
    
    doc.add_paragraph('numpy : Calculs numeriques et matriciels', style='List Bullet')
    doc.add_paragraph('pandas : Manipulation de donnees tabulaires', style='List Bullet')
    doc.add_paragraph('pulp : Interface de programmation lineaire', style='List Bullet')
    doc.add_paragraph('openpyxl : Generation fichiers Excel', style='List Bullet')
    doc.add_paragraph('python-docx : Generation rapports Word', style='List Bullet')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Solveur d\'optimisation :').bold = True
    p.add_run(' CBC (inclus automatiquement avec PuLP)')
    
    doc.add_paragraph()
    
    # Mise en production
    add_red_heading(doc, '5. MISE EN PRODUCTION', level=1)
    
    doc.add_heading('5.1 Pipeline de production quotidien', level=2)
    
    # Tableau pipeline production
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid'
    
    # Headers
    headers_row = table.rows[0]
    set_cell_background(headers_row.cells[0], 'DC143C')
    set_cell_background(headers_row.cells[1], 'DC143C')
    set_cell_background(headers_row.cells[2], 'DC143C')
    headers_row.cells[0].text = 'Phase'
    headers_row.cells[1].text = 'Actions'
    headers_row.cells[2].text = 'Moment'
    for cell in headers_row.cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Donnees
    prod_pipeline = [
        ('1. Collecte', 'Recuperation soldes caisse veille\nRecuperation predictions ML', 'Fin de journee'),
        ('2. Execution', 'Lancement automatique script\nGeneration fichiers Excel + Word', 'Planifie (nuit)'),
        ('3. Validation', 'Verification coherence\nAjustements manuels si necessaire', 'Debut matinee'),
        ('4. Distribution', 'Envoi plans d\'action aux agences\nSuivi execution transferts', 'Matinee')
    ]
    
    for i, (phase, actions, moment) in enumerate(prod_pipeline, 1):
        table.rows[i].cells[0].text = phase
        table.rows[i].cells[1].text = actions
        table.rows[i].cells[2].text = moment
    
    doc.add_paragraph()
    
    # Plan deploiement
    doc.add_heading('5.2 Plan de deploiement (3 phases)', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Phase 1 : Pilote (2 semaines)').bold = True
    
    doc.add_paragraph('Deploiement sur 5 agences test', style='List Bullet')
    doc.add_paragraph('Execution parallele du processus manuel', style='List Bullet')
    doc.add_paragraph('Validation des resultats et ajustements', style='List Bullet')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Phase 2 : Extension (1 mois)').bold = True
    
    doc.add_paragraph('Deploiement sur les 29 agences', style='List Bullet')
    doc.add_paragraph('Formation des equipes operationnelles', style='List Bullet')
    doc.add_paragraph('Monitoring renforce et optimisation parametres', style='List Bullet')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Phase 3 : Production complete (continu)').bold = True
    
    doc.add_paragraph('Automatisation totale du processus', style='List Bullet')
    doc.add_paragraph('Supervision quotidienne', style='List Bullet')
    doc.add_paragraph('Amelioration continue', style='List Bullet')
    
    doc.add_paragraph()
    
    # KPI
    doc.add_heading('5.3 Indicateurs de performance (KPI)', level=2)
    
    doc.add_paragraph('Cout moyen quotidien (FCFA)', style='List Number')
    doc.add_paragraph('Nombre de transactions par jour', style='List Number')
    doc.add_paragraph('Taux d\'intervention BEAC (%)', style='List Number')
    doc.add_paragraph('Temps de traitement (minutes)', style='List Number')
    doc.add_paragraph('Taux d\'erreur de prevision (MAPE %)', style='List Number')
    doc.add_paragraph('Economies realisees vs processus manuel', style='List Number')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Conclusion
    add_red_heading(doc, '6. CONCLUSION', level=1)
    
    doc.add_paragraph(
        'Le systeme d\'optimisation des flux de tresorerie par programmation lineaire '
        'offre une solution robuste, rapide et mathematiquement optimale. L\'utilisation '
        'du solveur CBC garantit la meilleure allocation possible des liquidites tout en '
        'minimisant les couts operationnels.'
    )
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'Les avantages cles incluent l\'optimalite garantie, la rapidite d\'execution, '
        'la robustesse face aux contraintes complexes, et la scalabilite pour des deployements '
        'futurs sur un nombre d\'agences etendu.'
    )
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'Le plan de deploiement en trois phases assure une transition progressive et securisee '
        'vers la production complete, avec validation a chaque etape et formation adequate '
        'des equipes operationnelles.'
    )
    
    # Sauvegarder
    filename = 'Rapport_Presentation_Optimisation_Tresorerie.docx'
    doc.save(filename)
    
    print("Rapport genere avec succes !")
    print(f"Fichier : {filename}")
    print("Pages : 5")
    print("Format : Professionnel sans emojis")
    print("Couleurs : Rouge et blanc")


if __name__ == "__main__":
    generate_presentation_report()
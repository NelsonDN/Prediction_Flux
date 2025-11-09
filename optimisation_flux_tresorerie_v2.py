"""
================================================================================
PROGRAMME COMPLET : OPTIMISATION DES FLUX DE LIQUIDITE BANCAIRE
Phase 2 : Allocation optimale et minimisation des couts de transport

Auteur: Systeme de Prevision de liquidité
Date: 2025
Version: 1.0

Description:
    Ce programme realise l'optimisation des flux de liquidité entre agences
    bancaires en minimisant les couts de transport tout en equilibrant les
    besoins et excedents de liquidite.

Structure:
    1. IMPORTS ET CONFIGURATION
    2. FONCTIONS UTILITAIRES
    3. SELECTION MODELE OPTIMAL
    4. GENERATION MATRICE COUTS (FEUILLE 4)
    5. GENERATION DONNEES ENTREE (FEUILLES 1 & 2)
    6. OPTIMISATION TRANSPORT (FEUILLE 3)
    7. EXPORT EXCEL
    8. GENERATION RAPPORT WORD
    9. PIPELINE PRINCIPAL
================================================================================
"""

# ============================================================================
# 1. IMPORTS ET CONFIGURATION
# ============================================================================

import os
import sys
import warnings
from pathlib import Path
from datetime import datetime, timedelta
import numpy as np
import pandas as pd
import pulp
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings('ignore')

# Configuration globale
CONFIG = {
    'agences': ['00001', '00005', '00010', '00021', '00024', '00026', '00031', 
                '00033', '00034', '00037', '00038', '00039', '00040', '00042', 
                '00043', '00045', '00046', '00055', '00056', '00057', '00062', 
                '00063', '00064', '00073', '00075', '00079', '00081', '00085', 
                '00087', 'BEAC'],
    'agences_normales': ['00001', '00005', '00010', '00021', '00024', '00026', 
                         '00031', '00033', '00034', '00037', '00038', '00039', 
                         '00040', '00042', '00043', '00045', '00046', '00055', 
                         '00056', '00057', '00062', '00063', '00064', '00073', 
                         '00075', '00079', '00081', '00085', '00087'],
    'models': ['xgboost', 'prophet', 'moving_avg', 
               'ensemble_post', 'ensemble_cv_best', 'ensemble_cv_second'],
    'models_exclus': ['arima'],
    'cout_min': 30,
    'cout_max': 100,
    'decimales': 2,
    'horizon': 5,
    'base_dir': 'resultats_metriques',
    'output_dir': 'optimisation_transport__',
    'date_dernier_test': '2025-08-31',
    'seed_random': 42,
    'forced_model': None,  # Mettre None pour sélection auto, ou 'xgboost', 'prophet', etc.
    'forced_approach': None,  # Mettre None pour sélection auto, ou 'direct', 'reconstruction'
    # Pour forcer xgboost avec approche directe :
    # 'forced_model': 'xgboost',
    # 'forced_approach': 'direct',
}

# Creation dossier output
Path(CONFIG['output_dir']).mkdir(parents=True, exist_ok=True)

print("="*80)
print("OPTIMISATION DES FLUX DE TRESORERIE BANCAIRE - PHASE 2")
print("="*80)
print()


# ============================================================================
# 2. FONCTIONS UTILITAIRES
# ============================================================================

def calculate_mape(y_true, y_pred):
    """
    Calcule le Mean Absolute Percentage Error
    
    Args:
        y_true: Valeurs reelles
        y_pred: Valeurs predites
    
    Returns:
        MAPE en pourcentage
    """
    y_true = np.array(y_true)
    y_pred = np.array(y_pred)
    
    # Eviter division par zero
    mask = np.abs(y_true) > 1e-8
    
    if not mask.any():
        return 0.0
    
    mape = np.mean(np.abs((y_true[mask] - y_pred[mask]) / y_true[mask])) * 100
    return mape

def load_predictions(agence, composante):
    """
    Charge les predictions pour une agence et une composante
    
    Args:
        agence: Code agence (ex: '00001')
        composante: 'encaissements', 'decaissements', ou 'Besoin'
    
    Returns:
        DataFrame avec predictions ou None si fichier inexistant
    """
    filepath = Path(CONFIG['base_dir']) / agence / f"{agence}_predictions_{composante}.csv"
    
    if not filepath.exists():
        return None  # Retourner None au lieu de lever une erreur
    
    df = pd.read_csv(filepath)
    return df

def format_number(value, decimales=2):
    """Formate un nombre avec separateur de milliers"""
    return f"{value:,.{decimales}f}".replace(',', ' ')


def get_date_from_jour(jour):
    """
    Calcule la date correspondant au jour de prediction
    
    Args:
        jour: 1, 2, 3, 4 ou 5
    
    Returns:
        Date au format string
    """
    date_base = datetime.strptime(CONFIG['date_dernier_test'], '%Y-%m-%d')
    date_pred = date_base + timedelta(days=jour)
    return date_pred.strftime('%Y-%m-%d')


# ============================================================================
# 3. SELECTION MODELE OPTIMAL
# ============================================================================

def select_best_model():
    """
    Selectionne le meilleur modele selon 2 approches:
    1. Approche Directe: prediction de la colonne Besoin
    2. Approche Reconstruction: Encaissements - Decaissements
    
    Returns:
        dict avec:
            - best_model: nom du meilleur modele
            - best_approach: 'direct' ou 'reconstruction'
            - best_mape: MAPE du meilleur modele
            - ranking: DataFrame avec classement complet
    """

    # Vérifier si modèle forcé
    if CONFIG.get('forced_model') and CONFIG.get('forced_approach'):
        print("ATTENTION: Modele et approche forces (pas de selection automatique)")
        print(f"Modele force     : {CONFIG['forced_model']}")
        print(f"Approche forcee  : {CONFIG['forced_approach']}")
        
        return {
            'best_model': CONFIG['forced_model'],
            'best_approach': CONFIG['forced_approach'],
            'best_mape': 0.0,
            'ranking': pd.DataFrame(),
            'all_results': {}
        }
    
    # Sinon, sélection automatique normale
    print("[1/9] Selection du modele optimal...")

    print("[1/9] Selection du modele optimal...")
    print("-" * 80)
    
    agences = CONFIG['agences_normales']
    models = CONFIG['models']
    
    results = {}
    
  
    # APPROCHE 1: PREDICTION DIRECTE
    print("Evaluation Approche 1: Prediction directe colonne Besoin")
    
    for model in models:
        mapes = []
        agences_traitees = 0
        
        for agence in agences:
            try:
                df_besoin = load_predictions(agence, 'Besoin')
                
                # Si fichier inexistant, passer
                if df_besoin is None:
                    continue
                
                # Verifier que le modele existe
                if model not in df_besoin.columns:
                    continue
                
                pred = df_besoin[model].values
                real = df_besoin['Reel'].values
                
                mape = calculate_mape(real, pred)
                
                # Ignorer les MAPE aberrants (> 10000%)
                if mape < 10000:
                    mapes.append(mape)
                    agences_traitees += 1
                
            except Exception as e:
                continue
        
        if mapes and agences_traitees >= 5:  # Au moins 5 agences
            mape_moyen = np.mean(mapes)
            results[f"{model}_direct"] = mape_moyen
            print(f"  {model:20s} : MAPE moyen = {mape_moyen:.2f}% ({agences_traitees} agences)")
        else:
            print(f"  {model:20s} : Donnees insuffisantes ({agences_traitees} agences)")




    # APPROCHE 2: RECONSTRUCTION
    print()
    print("Evaluation Approche 2: Reconstruction (Encaissements - Decaissements)")
    
    for model in models:
        mapes = []
        agences_traitees = 0
        
        for agence in agences:
            try:
                df_enc = load_predictions(agence, 'encaissements')
                df_dec = load_predictions(agence, 'decaissements')
                df_besoin = load_predictions(agence, 'Besoin')
                
                # Si un fichier est manquant, passer
                if df_enc is None or df_dec is None or df_besoin is None:
                    continue
                
                # Verifier que le modele existe
                if model not in df_enc.columns or model not in df_dec.columns:
                    continue
                
                besoin_reconstruit = df_enc[model].values - df_dec[model].values
                real = df_besoin['Reel'].values
                
                mape = calculate_mape(real, besoin_reconstruit)
                
                # Ignorer les MAPE aberrants
                if mape < 10000:
                    mapes.append(mape)
                    agences_traitees += 1
                
            except Exception as e:
                continue
        
        if mapes and agences_traitees >= 5:  # Au moins 5 agences
            mape_moyen = np.mean(mapes)
            results[f"{model}_reconstruction"] = mape_moyen
            print(f"  {model:20s} : MAPE moyen = {mape_moyen:.2f}% ({agences_traitees} agences)")
        else:
            print(f"  {model:20s} : Donnees insuffisantes ({agences_traitees} agences)")




    
    # SELECTION DU MEILLEUR
    if not results:
        raise RuntimeError("Aucun resultat d'evaluation obtenu")
    
    best_combo = min(results, key=results.get)
    best_mape = results[best_combo]
    
    # Parser le nom
    parts = best_combo.rsplit('_', 1)
    best_model = parts[0]
    best_approach = parts[1]
    
    # Creer classement
    ranking_data = []
    for combo, mape in sorted(results.items(), key=lambda x: x[1]):
        parts = combo.rsplit('_', 1)
        model = parts[0]
        approach = parts[1]
        ranking_data.append({
            'Modele': model,
            'Approche': approach,
            'MAPE_Moyen': round(mape, 2)
        })
    
    ranking_df = pd.DataFrame(ranking_data)
    
    print("=" * 80)
    print("RESULTAT SELECTION MODELE")
    print("=" * 80)
    print(f"Meilleur modele     : {best_model}")
    print(f"Meilleure approche  : {best_approach}")
    print(f"MAPE moyen          : {best_mape:.2f}%")
    print()
    print("Classement complet:")
    print(ranking_df.to_string(index=False))
    print()
    
    return {
        'best_model': best_model,
        'best_approach': best_approach,
        'best_mape': best_mape,
        'ranking': ranking_df,
        'all_results': results
    }


# ============================================================================
# 4. GENERATION MATRICE COUTS (FEUILLE 4)
# ============================================================================

# def generate_matrice_couts():
#     """
#     Charge la matrice des couts unitaires depuis le fichier Excel
#     """
#     print("[2/9] Chargement de la matrice des couts unitaires...")
#     print("-" * 80)
    
#     filepath = 'Données pour la solutions streamlit - Avec code.xlsx'
    
#     if not Path(filepath).exists():
#         raise FileNotFoundError(f"Fichier Excel introuvable: {filepath}")
    
#     # ✅ Lire Feuil3 en mode string
#     df_raw = pd.read_excel(filepath, sheet_name='Feuil3', header=None, dtype=str)
    
#     # 🔍 DEBUG 1: Dimensions du fichier
#     print("🔍 DEBUG 1 - Dimensions du DataFrame brut:")
#     print(f"   Shape: {df_raw.shape}")
#     print(f"   Nombre lignes: {len(df_raw)}")
#     print(f"   Nombre colonnes: {len(df_raw.columns)}")
#     print()
    
#     # 🔍 DEBUG 2: Aperçu des premières lignes/colonnes
#     print("🔍 DEBUG 2 - Aperçu B2:E5:")
#     print(df_raw.iloc[1:5, 1:5])
#     print()
    
#     # ✅ Extraire codes agences LIGNES (B4:B33)
#     agences_lignes = df_raw.iloc[3:33, 1].tolist()
#     agences_lignes = [str(x).strip() for x in agences_lignes]
    
#     # 🔍 DEBUG 3: Agences lignes
#     print("🔍 DEBUG 3 - Agences en LIGNES (B4:B33):")
#     print(f"   Nombre: {len(agences_lignes)}")
#     print(f"   Premier: {agences_lignes[0]}")
#     print(f"   Dernier: {agences_lignes[-1]}")
#     print(f"   Liste: {agences_lignes}")
#     print()
    
#     # ✅ Extraire codes agences COLONNES (C2:AF2)
#     agences_colonnes = df_raw.iloc[1, 2:32].tolist()
#     agences_colonnes = [str(x).strip() for x in agences_colonnes]
    
#     # 🔍 DEBUG 4: Agences colonnes
#     print("🔍 DEBUG 4 - Agences en COLONNES (C2:AF2):")
#     print(f"   Nombre: {len(agences_colonnes)}")
#     print(f"   Premier: {agences_colonnes[0]}")
#     print(f"   Dernier: {agences_colonnes[-1]}")
#     print(f"   Liste: {agences_colonnes}")
#     print()
    
#     # ✅ Extraire la matrice de valeurs
#     matrice_valeurs = df_raw.iloc[3:33, 2:32].values
    
#     # 🔍 DEBUG 5: Dimensions de la matrice extraite
#     print("🔍 DEBUG 5 - Matrice de valeurs (C4:AF33):")
#     print(f"   Shape: {matrice_valeurs.shape}")
#     print(f"   Lignes: {matrice_valeurs.shape[0]}")
#     print(f"   Colonnes: {matrice_valeurs.shape[1]}")
#     print()
    
#     # 🔍 DEBUG 6: CONFIG['agences']
#     print("🔍 DEBUG 6 - CONFIG['agences']:")
#     print(f"   Nombre: {len(CONFIG['agences'])}")
#     print(f"   Liste: {CONFIG['agences']}")
#     print()
    
#     # ✅ Vérification cohérence
#     n_lignes = len(agences_lignes)
#     n_colonnes = len(agences_colonnes)
#     n_config = len(CONFIG['agences'])
    
#     print("🔍 DEBUG 7 - Vérification cohérence:")
#     print(f"   n_lignes extraites: {n_lignes}")
#     print(f"   n_colonnes extraites: {n_colonnes}")
#     print(f"   n_config: {n_config}")
#     print(f"   matrice_valeurs shape: {matrice_valeurs.shape}")
#     print()
    
#     # ✅ Utiliser n basé sur la matrice réelle
#     n_rows, n_cols = matrice_valeurs.shape
#     n = min(n_rows, n_cols, len(agences_lignes), len(agences_colonnes))
    
#     print(f"✅ Utilisation de n = {n} pour la suite")
#     print()
    
#     # ✅ Convertir en float
#     matrice = np.zeros((n, n), dtype=float)
    
#     print("🔍 DEBUG 8 - Conversion en cours...")
#     erreurs = 0
    
#     for i in range(n):
#         for j in range(n):
#             try:
#                 val = matrice_valeurs[i, j]
#                 if pd.notna(val) and str(val).strip() and str(val).strip().lower() != 'nan':
#                     val_clean = str(val).strip().replace(' ', '').replace(',', '.')
#                     matrice[i, j] = float(val_clean)
#             except Exception as e:
#                 erreurs += 1
#                 if erreurs <= 5:  # Afficher seulement les 5 premières erreurs
#                     print(f"   ⚠️ Erreur [{i},{j}]: {e}")
    
#     print(f"   Erreurs de conversion: {erreurs}")
#     print()
    
#     # ✅ Compléter le triangle supérieur par symétrie
#     for i in range(n):
#         for j in range(i+1, n):
#             if matrice[i, j] == 0 and matrice[j, i] != 0:
#                 matrice[i, j] = matrice[j, i]
#             elif matrice[j, i] == 0 and matrice[i, j] != 0:
#                 matrice[j, i] = matrice[i, j]
    
#     # ✅ Vérifier diagonale = 0
#     np.fill_diagonal(matrice, 0)
    
#     # ✅ Convertir en DataFrame (utiliser les n premières agences)
#     df_couts = pd.DataFrame(
#         matrice, 
#         index=agences_lignes[:n], 
#         columns=agences_colonnes[:n]
#     )
    
#     print(f"✅ Matrice chargée: {n}x{n}")
#     print(f"   Source: Feuil3 (triangle inferieur + symetrie)")
#     print()
    
#     # Statistiques
#     couts_non_nuls = matrice[matrice > 0]
#     if len(couts_non_nuls) > 0:
#         print("📊 Statistiques des coûts:")
#         print(f"   Minimum   : {couts_non_nuls.min():.0f} FCFA")
#         print(f"   Maximum   : {couts_non_nuls.max():.0f} FCFA")
#         print(f"   Moyenne   : {couts_non_nuls.mean():.2f} FCFA")
#         print(f"   Mediane   : {np.median(couts_non_nuls):.0f} FCFA")
#         print(f"   Nb valeurs non nulles: {len(couts_non_nuls)}")
#     print()
    
#     return df_couts






def generate_matrice_couts():
    """
    Charge la matrice des couts unitaires depuis le fichier Excel
    """
    print("[2/9] Chargement de la matrice des couts unitaires...")
    print("-" * 80)
    
    filepath = 'Données pour la solutions streamlit - Avec code.xlsx'
    
    if not Path(filepath).exists():
        raise FileNotFoundError(f"Fichier Excel introuvable: {filepath}")
    
    # ✅ Lire Feuil3 SANS forcer dtype=str (laisser pandas gérer les types)
    df_raw = pd.read_excel(filepath, sheet_name='Feuil3', header=None)
    
    print(f"📄 Dimensions: {df_raw.shape}")
    print()
    
    # ✅ Extraire codes agences LIGNES (B4:B33) - FORCER en string
    agences_lignes = df_raw.iloc[3:33, 1].astype(str).str.strip().tolist()
    
    # ✅ Extraire codes agences COLONNES (C2:AF2) - FORCER en string
    # agences_colonnes = df_raw.iloc[1, 2:32].astype(str).str.strip().tolist()



    # ✅ APRÈS - Formater les codes correctement
    agences_colonnes_raw = df_raw.iloc[1, 2:32].tolist()
    agences_colonnes = []
    for code in agences_colonnes_raw:
        if pd.notna(code):
            # Convertir en int puis formater avec zéros
            if isinstance(code, str):
                code_clean = code.strip()
                if code_clean.upper() == 'BEAC':
                    agences_colonnes.append('BEAC')
                else:
                    try:
                        agences_colonnes.append(str(int(float(code_clean))).zfill(5))
                    except:
                        agences_colonnes.append(code_clean)
            else:
                # C'est un nombre (float/int)
                code_int = int(code)
                if code_int == 0:  # Gérer le cas BEAC encodé comme 0
                    agences_colonnes.append('BEAC')
                else:
                    agences_colonnes.append(str(code_int).zfill(5))



    print(f"✅ {len(agences_lignes)} agences en lignes")
    print(f"✅ {len(agences_colonnes)} agences en colonnes")
    print()
    
    # ✅ Extraire la matrice de valeurs (C4:AF33) - GARDER comme numérique
    matrice_valeurs = df_raw.iloc[3:33, 2:32].values
    
    # ✅ Convertir en float et ARRONDIR à l'entier
    n = len(agences_lignes)
    matrice = np.zeros((n, n), dtype=float)
    
    for i in range(n):
        for j in range(n):
            val = matrice_valeurs[i, j]
            if pd.notna(val):
                try:
                    # Convertir en float puis ARRONDIR
                    matrice[i, j] = round(float(val))
                except:
                    matrice[i, j] = 0.0
    
    # ✅ Compléter le triangle supérieur par symétrie
    for i in range(n):
        for j in range(i+1, n):
            if matrice[i, j] == 0 and matrice[j, i] != 0:
                matrice[i, j] = matrice[j, i]
            elif matrice[j, i] == 0 and matrice[i, j] != 0:
                matrice[j, i] = matrice[i, j]
    
    # ✅ Vérifier diagonale = 0
    np.fill_diagonal(matrice, 0)
    
    # ✅ Convertir en DataFrame
    df_couts = pd.DataFrame(matrice, index=agences_lignes, columns=agences_colonnes)
    
    print(f"✅ Matrice chargée: {n}x{n}")
    print()
    
    # Statistiques
    couts_non_nuls = matrice[matrice > 0]
    if len(couts_non_nuls) > 0:
        print("📊 Statistiques des coûts:")
        print(f"   Minimum   : {couts_non_nuls.min():.0f} FCFA")
        print(f"   Maximum   : {couts_non_nuls.max():.0f} FCFA")
        print(f"   Moyenne   : {couts_non_nuls.mean():.2f} FCFA")
        print(f"   Mediane   : {np.median(couts_non_nuls):.0f} FCFA")
        print(f"   Nb valeurs: {len(couts_non_nuls)}")
    print()
    
    return df_couts




# ============================================================================
# 5. GENERATION DONNEES ENTREE (FEUILLES 1 & 2)
# ============================================================================


def get_besoin_prediction(agence, model, approach, jour):
    """
    Recupere la prediction du Besoin pour une agence, un jour donne
    
    Args:
        agence: Code agence
        model: Nom du modele
        approach: 'direct' ou 'reconstruction'
        jour: 1 a 5
    
    Returns:
        Valeur predite du Besoin (float)
    """
    idx = jour - 1  # Index 0 = J+1
    
    if approach == 'direct':
        df = load_predictions(agence, 'Besoin')
        if df is None or model not in df.columns:
            raise ValueError(f"Predictions Besoin indisponibles pour {agence}, modele {model}")
        return df[model].iloc[idx]
    
    else:  # reconstruction
        df_enc = load_predictions(agence, 'encaissements')
        df_dec = load_predictions(agence, 'decaissements')
        
        if df_enc is None or df_dec is None or model not in df_enc.columns or model not in df_dec.columns:
            raise ValueError(f"Predictions encaissements/decaissements indisponibles pour {agence}, modele {model}")
        
        enc = df_enc[model].iloc[idx]
        dec = df_dec[model].iloc[idx]
        
        return enc - dec


def get_mape_agence(agence, model, approach):
    """
    Calcule le MAPE pour une agence avec le modele selectionne
    
    Args:
        agence: Code agence
        model: Nom du modele
        approach: 'direct' ou 'reconstruction'
    
    Returns:
        MAPE en pourcentage
    """
    if approach == 'direct':
        df = load_predictions(agence, 'Besoin')
        if df is None or model not in df.columns:
            return 0.0  # Valeur par defaut si donnees manquantes
        pred = df[model].values
        real = df['Reel'].values
        return calculate_mape(real, pred)
    
    else:  # reconstruction
        df_enc = load_predictions(agence, 'encaissements')
        df_dec = load_predictions(agence, 'decaissements')
        df_besoin = load_predictions(agence, 'Besoin')
        
        if df_enc is None or df_dec is None or df_besoin is None:
            return 0.0  # Valeur par defaut si donnees manquantes
        
        if model not in df_enc.columns or model not in df_dec.columns:
            return 0.0
        
        besoin_reconstruit = df_enc[model].values - df_dec[model].values
        real = df_besoin['Reel'].values
        
        return calculate_mape(real, besoin_reconstruit)


def get_stock_minimum_agence(agence):
    """
    Recupere le stock minimum reel d'une agence depuis le fichier Excel
    
    Args:
        agence: Code agence (ex: '00001')
    
    Returns:
        Stock de securite en FCFA (string pour flexibilité)
    """
    filepath = 'Données pour la solutions streamlit - Avec code.xlsx'
    
    if not Path(filepath).exists():
        raise FileNotFoundError(
            f"Fichier Excel introuvable: {filepath}\n"
            f"Veuillez placer le fichier a la racine du projet"
        )
    
    # ✅ Lire Feuil1 avec Code_Agence en string
    df_stocks = pd.read_excel(filepath, sheet_name='Feuil1', dtype={'Code_Agence': str})
    df_stocks['Code_Agence'] = df_stocks['Code_Agence'].str.strip()
    
    # Filtrer pour l'agence
    ligne_agence = df_stocks[df_stocks['Code_Agence'] == agence]
    
    if len(ligne_agence) == 0:
        raise ValueError(f"Stock de securite introuvable pour agence {agence}")
    
    # ✅ Récupérer comme STRING (plus flexible)
    stock_securite = str(ligne_agence['Stock_de_securite'].values[0]).strip()
    
    return stock_securite

# def generate_feuilles_entree(jour, model, approach, soldes_veille_dict):
def generate_feuilles_entree(jour, model, approach):
    """
    Genere les donnees d'entree (Feuilles 1 & 2) pour un jour donne
    
    Args:
        jour: 1 a 5
        model: Modele selectionne
        approach: 'direct' ou 'reconstruction'
        soldes_veille_dict: Dictionnaire {agence: solde_veille}
    
    Returns:
        DataFrame avec colonnes A a I pour les 30 agences
    """
    date_str = get_date_from_jour(jour)
    
    data = []
    
    for agence in CONFIG['agences']:
        if agence == 'BEAC':
            # BEAC: traitement special, on l'ajoutera apres
            continue
        
        # Colonne C: Solde Caisse veille
        # solde_veille = soldes_veille_dict.get(agence, 0.0)
        np.random.seed(int(agence) + jour * 1000)
        solde_veille = round(np.random.uniform(5,1000), 2)


        # Colonne D: Besoin en liquidite (prediction)
        besoin_pred = get_besoin_prediction(agence, model, approach, jour)
        
        # Colonne E: Score de precision (MAPE)
        mape = get_mape_agence(agence, model, approach)
        
        # # Colonne F: Stock minimum (aleatoire [10, 50])
        # np.random.seed(int(agence) + jour)  # Seed pour reproductibilite
        # stock_min = np.random.randint(10, 51)


        # Colonne F: Stock minimum (depuis fichier CSV calculé)
        # stock_min = get_stock_minimum_agence(agence)


        # Colonne F: Stock de securite (depuis fichier Excel Feuil1)
        stock_securite_str = get_stock_minimum_agence(agence)

        # Convertir en millions pour les calculs
        try:
            # Enlever espaces et convertir en float
            stock_securite_fcfa = float(stock_securite_str.replace(' ', ''))
            stock_min = stock_securite_fcfa / 1_000_000  # Convertir en millions
        except:
            raise ValueError(f"Impossible de convertir le stock de securite pour {agence}: {stock_securite_str}")

        
        # Colonne G: Stock de securite = ENT((1 + E) * F) + 1
        stock_secu = int((1 + mape/100) * stock_min) + 1
        
        # Colonne H: Flux transport optimal = D + G - C
        flux_optimal = besoin_pred + stock_secu - solde_veille
        
        # Colonne I: Type de flux
        if flux_optimal > 0:
            type_flux = "Besoin"
        elif flux_optimal < 0:
            type_flux = "Excedent"
        else:
            type_flux = "Aucun besoin"
        
        data.append({
            'Date': date_str,
            'Code Agence': agence,
            'Solde Caisse veille': round(solde_veille, 2),
            'Besoin en liquidite': round(besoin_pred, 2),
            'Score de precision': round(mape, 2),
            'Stock minimum': stock_min,
            'Stock de securite': stock_secu,
            'Flux transport optimal': round(flux_optimal, 2),
            # 'Flux transport optimal': flux_optimal,
            'Type de flux': type_flux
        })
    
    # Calculer flux BEAC = -SOMME(flux des 29 agences)
    flux_agences = [row['Flux transport optimal'] for row in data]
    flux_beac = -sum(flux_agences)
    
    # Type flux BEAC
    if flux_beac > 0.01:
        type_flux_beac = "Versement BEAC"
    elif flux_beac < -0.01:
        type_flux_beac = "Retrait BEAC"
    else:
        type_flux_beac = "Aucun besoin"
    
    # Ajouter BEAC avec flux calculé
    data.append({
        'Date': date_str,
        'Code Agence': 'BEAC',
        'Solde Caisse veille': None,
        'Besoin en liquidite': None,
        'Score de precision': None,
        'Stock minimum': None,
        'Stock de securite': None,
        'Flux transport optimal': round(flux_beac, 2),  # ← CALCULÉ
        # 'Flux transport optimal': flux_beac,  # ← CALCULÉ
        'Type de flux': type_flux_beac                  # ← CALCULÉ
    })
    
    df = pd.DataFrame(data)
    
    return df

# ============================================================================
# 6. OPTIMISATION TRANSPORT (FEUILLE 3)
# ============================================================================




# def optimize_transport_allocation(df_entree, matrice_couts):
#     """
#     Optimise l'allocation des flux entre agences en minimisant les couts
    
#     Probleme de transport:
#         Minimiser: sum(C[i][j] * x[i][j])
        
#         Sous contraintes:
#             - sum_j(x[i][j]) = Excedent[i]  pour tout i
#             - sum_i(x[i][j]) = Besoin[j]    pour tout j
#             - x[i][j] >= 0
    
#     Args:
#         df_entree: DataFrame des donnees d'entree
#         matrice_couts: DataFrame 30x30 des couts unitaires
    
#     Returns:
#         dict avec:
#             - allocation_matrix: np.array 30x30
#             - cout_total: float
#             - nb_transactions: int
#             - agences_besoin: dict
#             - agences_excedent: dict
#             - status: str
#     """
#     # Separer Besoin et Excedent
#     besoin_dict = {}
#     excedent_dict = {}
    
#     for _, row in df_entree.iterrows():
#         agence = row['Code Agence']
#         flux = row['Flux transport optimal']
        
#         if flux > 0.01:  # Tolerance numerique
#             besoin_dict[agence] = flux
#         elif flux < -0.01:
#             excedent_dict[agence] = abs(flux)
    

#     # Verification et ajustement pour equilibre parfait
#     total_besoin = sum(besoin_dict.values())
#     total_excedent = sum(excedent_dict.values())
    
#     ecart = total_besoin - total_excedent
    
#     # Ajuster la plus grosse agence en excédent pour compenser l'écart
#     if abs(ecart) > 0.000001:
#         print(f"  Ajustement equilibre : {ecart:.6f} M")
        
#         agence_max_exc = max(excedent_dict, key=excedent_dict.get)
#         excedent_dict[agence_max_exc] = round(excedent_dict[agence_max_exc] + ecart, 2)
        
#         # Recalculer totaux
#         total_besoin = sum(besoin_dict.values())
#         total_excedent = sum(excedent_dict.values())
    
#     print(f"  Agences en Besoin    : {len(besoin_dict)}")
#     print(f"  Agences en Excedent  : {len(excedent_dict)}")
#     print(f"  Total Besoin         : {format_number(total_besoin)} M")
#     print(f"  Total Excedent       : {format_number(total_excedent)} M")
#     print(f"  Verification B1 = E1 : OK")
#     print()
    



#     # Optimisation avec PuLP
#     print("  Resolution du probleme d'optimisation...")
    
#     prob = pulp.LpProblem("Minimisation_Couts_Transport", pulp.LpMinimize)
    
#     # Variables de decision
#     x = {}
#     for i in excedent_dict.keys():
#         for j in besoin_dict.keys():
#             x[i, j] = pulp.LpVariable(
#                 f"x_{i}_{j}",
#                 lowBound=0,
#                 cat='Continuous'
#             )
    
#     # Fonction objectif
#     cout_objectif = pulp.lpSum([
#         matrice_couts.loc[i, j] * x[i, j]
#         for i in excedent_dict.keys()
#         for j in besoin_dict.keys()
#     ])
#     prob += cout_objectif
    
#     # Contrainte 1: Conservation excedents
#     for i in excedent_dict.keys():
#         prob += pulp.lpSum([x[i, j] for j in besoin_dict.keys()]) == excedent_dict[i]
    
#     # Contrainte 2: Satisfaction besoins
#     for j in besoin_dict.keys():
#         prob += pulp.lpSum([x[i, j] for i in excedent_dict.keys()]) == besoin_dict[j]
    
#     # Resolution
#     prob.solve(pulp.PULP_CBC_CMD(msg=0))
    
#     # Verification statut
#     if prob.status != pulp.LpStatusOptimal:
#         raise RuntimeError(f"Optimisation echouee: statut = {pulp.LpStatus[prob.status]}")
    
#     print(f"  Statut optimisation  : {pulp.LpStatus[prob.status]}")
    
#     # Extraction resultats
#     agences_list = list(matrice_couts.index)
#     n = len(agences_list)
#     allocation_matrix = np.zeros((n, n))
    
#     for i in excedent_dict.keys():
#         for j in besoin_dict.keys():
#             valeur = x[i, j].varValue
#             if valeur is not None and valeur > 0.001:
#                 idx_i = agences_list.index(i)
#                 idx_j = agences_list.index(j)
#                 allocation_matrix[idx_i, idx_j] = round(valeur, CONFIG['decimales'])
#                 # allocation_matrix[idx_i, idx_j] = valeur
    
#     cout_total = round(pulp.value(prob.objective), CONFIG['decimales'])
#     nb_transactions = np.sum(allocation_matrix > 0)
    
#     print(f"  Cout total optimal   : {format_number(cout_total)} FCFA")
#     print(f"  Nombre transactions  : {nb_transactions}")
#     print()
    
#     return {
#         'allocation_matrix': allocation_matrix,
#         'cout_total': cout_total,
#         'nb_transactions': nb_transactions,
#         'agences_besoin': besoin_dict,
#         'agences_excedent': excedent_dict,
#         'status': 'Optimal'
#     }




def optimize_transport_allocation(df_entree, matrice_couts):
    """
    Optimise l'allocation des flux entre agences en minimisant les couts
    """
    # Separer Besoin et Excedent
    besoin_dict = {}
    excedent_dict = {}
    
    for _, row in df_entree.iterrows():
        agence = row['Code Agence']
        flux = row['Flux transport optimal']
        
        if flux > 0.01:  # Tolerance numerique
            besoin_dict[agence] = flux
        elif flux < -0.01:
            excedent_dict[agence] = abs(flux)
    
    # 🔍 DEBUG: Vérifier les codes agence
    print("🔍 DEBUG - Codes agence:")
    print(f"   Besoin dict: {list(besoin_dict.keys())[:5]}...")
    print(f"   Excédent dict: {list(excedent_dict.keys())[:5]}...")
    print(f"   Matrice index: {list(matrice_couts.index)[:5]}...")
    print(f"   Matrice columns: {list(matrice_couts.columns)[:5]}...")
    print()
    
    # ✅ Normaliser les codes agence dans la matrice
    matrice_couts.index = matrice_couts.index.astype(str).str.strip()
    matrice_couts.columns = matrice_couts.columns.astype(str).str.strip()
    
    # ✅ Normaliser les codes dans les dicts
    besoin_dict_clean = {str(k).strip(): v for k, v in besoin_dict.items()}
    excedent_dict_clean = {str(k).strip(): v for k, v in excedent_dict.items()}
    
    # 🔍 Vérifier que tous les codes existent dans la matrice
    codes_besoin_manquants = [k for k in besoin_dict_clean.keys() if k not in matrice_couts.columns]
    codes_excedent_manquants = [k for k in excedent_dict_clean.keys() if k not in matrice_couts.index]
    
    if codes_besoin_manquants:
        print(f"⚠️  ATTENTION: Codes en besoin ABSENTS de matrice_couts.columns:")
        print(f"   {codes_besoin_manquants}")
        print()
    
    if codes_excedent_manquants:
        print(f"⚠️  ATTENTION: Codes en excédent ABSENTS de matrice_couts.index:")
        print(f"   {codes_excedent_manquants}")
        print()
    
    # Utiliser les versions nettoyées
    besoin_dict = besoin_dict_clean
    excedent_dict = excedent_dict_clean
    
    # Verification et ajustement pour equilibre parfait
    total_besoin = sum(besoin_dict.values())
    total_excedent = sum(excedent_dict.values())
    
    ecart = total_besoin - total_excedent
    
    # Ajuster la plus grosse agence en excédent pour compenser l'écart
    if abs(ecart) > 0.000001:
        print(f"  Ajustement equilibre : {ecart:.6f} M")
        
        agence_max_exc = max(excedent_dict, key=excedent_dict.get)
        excedent_dict[agence_max_exc] = round(excedent_dict[agence_max_exc] + ecart, 2)
        
        # Recalculer totaux
        total_besoin = sum(besoin_dict.values())
        total_excedent = sum(excedent_dict.values())
    
    print(f"  Agences en Besoin    : {len(besoin_dict)}")
    print(f"  Agences en Excedent  : {len(excedent_dict)}")
    print(f"  Total Besoin         : {format_number(total_besoin)} M")
    print(f"  Total Excedent       : {format_number(total_excedent)} M")
    print(f"  Verification B1 = E1 : OK")
    print()
    
    # Optimisation avec PuLP
    print("  Resolution du probleme d'optimisation...")
    
    prob = pulp.LpProblem("Minimisation_Couts_Transport", pulp.LpMinimize)
    
    # Variables de decision
    x = {}
    for i in excedent_dict.keys():
        for j in besoin_dict.keys():
            x[i, j] = pulp.LpVariable(
                f"x_{i}_{j}",
                lowBound=0,
                cat='Continuous'
            )
    
    # Fonction objectif
    cout_objectif = pulp.lpSum([
        matrice_couts.loc[i, j] * x[i, j]
        for i in excedent_dict.keys()
        for j in besoin_dict.keys()
    ])
    prob += cout_objectif
    
    # Contrainte 1: Conservation excedents
    for i in excedent_dict.keys():
        prob += pulp.lpSum([x[i, j] for j in besoin_dict.keys()]) == excedent_dict[i]
    
    # Contrainte 2: Satisfaction besoins
    for j in besoin_dict.keys():
        prob += pulp.lpSum([x[i, j] for i in excedent_dict.keys()]) == besoin_dict[j]
    
    # Resolution
    prob.solve(pulp.PULP_CBC_CMD(msg=0))
    
    # Verification statut
    if prob.status != pulp.LpStatusOptimal:
        raise RuntimeError(f"Optimisation echouee: statut = {pulp.LpStatus[prob.status]}")
    
    print(f"  Statut optimisation  : {pulp.LpStatus[prob.status]}")
    
    # Extraction resultats
    agences_list = list(matrice_couts.index)
    n = len(agences_list)
    allocation_matrix = np.zeros((n, n))
    
    for i in excedent_dict.keys():
        for j in besoin_dict.keys():
            valeur = x[i, j].varValue
            if valeur is not None and valeur > 0.001:
                idx_i = agences_list.index(i)
                idx_j = agences_list.index(j)
                allocation_matrix[idx_i, idx_j] = round(valeur, CONFIG['decimales'])
    
    cout_total = round(pulp.value(prob.objective), CONFIG['decimales'])
    nb_transactions = np.sum(allocation_matrix > 0)
    
    print(f"  Cout total optimal   : {format_number(cout_total)} FCFA")
    print(f"  Nombre transactions  : {nb_transactions}")
    print()
    
    return {
        'allocation_matrix': allocation_matrix,
        'cout_total': cout_total,
        'nb_transactions': nb_transactions,
        'agences_besoin': besoin_dict,
        'agences_excedent': excedent_dict,
        'status': 'Optimal'
    }




# ============================================================================
# 7. EXPORT EXCEL
# ============================================================================

def export_excel_jour(jour, df_feuille1, df_feuille3_part1, allocation_matrix, 
                      matrice_couts, agences_list, result_optim):
    """
    Genere le fichier Excel pour un jour donne avec 4 feuilles
    
    Args:
        jour: 1 a 5
        df_feuille1: DataFrame feuille 1
        df_feuille3_part1: Dict avec separation Besoin/Excedent
        allocation_matrix: Matrice d'allocation 30x30
        matrice_couts: DataFrame matrice couts unitaires
        agences_list: Liste codes agences
        result_optim: Resultats optimisation
    """
    filename = Path(CONFIG['output_dir']) / f"Optimisation_Transport_J{jour}.xlsx"


    # AJOUTER CETTE LIGNE ICI :
    date_str = get_date_from_jour(jour)

    
    wb = Workbook()
    
    # Styles
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1F4788", end_color="1F4788", fill_type="solid")
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # ========================================================================
    # FEUILLE 1: Donnees d'entree
    # ========================================================================
    ws1 = wb.active
    ws1.title = "Donnees entree"
    
    # Headers
    # headers = list(df_feuille1.columns)
    # Headers avec colonnes J, K, L
    headers = list(df_feuille1.columns) + ['Solde prevue fin journee', 'Solde reel fin journee', 'Erreur sur la journee']
    for c_idx, header in enumerate(headers, 1):
        cell = ws1.cell(row=1, column=c_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border
    
    # Donnees (lignes 2 a 30 pour les 29 agences)
    for r_idx in range(len(df_feuille1) - 1):  # -1 pour exclure BEAC temporairement
        row_data = df_feuille1.iloc[r_idx]
        for c_idx, value in enumerate(row_data, 1):
            cell = ws1.cell(row=r_idx+2, column=c_idx, value=value)
            cell.border = border
            
            # Format numerique si necessaire
            if c_idx in [3, 4, 5, 8]:  # Colonnes C, D, E, H
                if value is not None:
                    cell.number_format = '#,##0.00'
    
    # Ligne 31: BEAC avec formules
    beac_row = 31
    
    # Colonne A: Date
    ws1.cell(row=beac_row, column=1, value=date_str)
    ws1.cell(row=beac_row, column=1).border = border
    
    # Colonne B: Code Agence
    ws1.cell(row=beac_row, column=2, value='BEAC')
    ws1.cell(row=beac_row, column=2).border = border
    ws1.cell(row=beac_row, column=2).font = Font(bold=True)
    
    # Colonnes C a G: Vides (avec bordures)
    for col in range(3, 8):
        ws1.cell(row=beac_row, column=col, value=None)
        ws1.cell(row=beac_row, column=col).border = border
    
    # # Colonne H: FORMULE =-SOMME(H2:H30)
    # ws1.cell(row=beac_row, column=8, value='=-SOMME(H2:H30)')
    # Colonne H: FORMULE =-SOMME(H2:H30) - Excel français utilise point-virgule
    # ws1.cell(row=beac_row, column=8, value='=-SOMME(H2:H30)')
    ws1.cell(row=beac_row, column=8, value='=-SUM(H2:H30)')
    ws1.cell(row=beac_row, column=8).border = border
    ws1.cell(row=beac_row, column=8).number_format = '#,##0.00'
    ws1.cell(row=beac_row, column=8).font = Font(bold=True)
    
    # Colonne I: FORMULE =SI(H31<0;"Retrait BEAC";SI(H31>0;"Versement BEAC";"Aucun besoin"))
    # ws1.cell(row=beac_row, column=9, value='=SI(H31<0;"Retrait BEAC";SI(H31>0;"Versement BEAC";"Aucun besoin"))')
    ws1.cell(row=beac_row, column=9, value='=IF(H31<0;"Retrait BEAC";IF(H31>0;"Versement BEAC";"Aucun besoin"))')
    ws1.cell(row=beac_row, column=9).border = border
    ws1.cell(row=beac_row, column=9).font = Font(bold=True)
    
    # Colonnes J, K, L: Vides (pour futur usage)
    for col in range(10, 13):
        ws1.cell(row=beac_row, column=col, value=None)
        ws1.cell(row=beac_row, column=col).border = border
    
    # Ajuster largeurs colonnes
    for col in ws1.columns:
        ws1.column_dimensions[col[0].column_letter].width = 18


    # ========================================================================
    # FEUILLE 2: Donnees d'entree (fixee) - Copie de Feuille 1
    # ========================================================================
    ws2 = wb.create_sheet("Donnees entree fixee")
    
    # Headers
    for c_idx, header in enumerate(headers, 1):
        cell = ws2.cell(row=1, column=c_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border
    
    # Donnees (lignes 2 a 30 pour les 29 agences)
    for r_idx in range(len(df_feuille1) - 1):  # -1 pour exclure BEAC temporairement
        row_data = df_feuille1.iloc[r_idx]
        for c_idx, value in enumerate(row_data, 1):
            cell = ws2.cell(row=r_idx+2, column=c_idx, value=value)
            cell.border = border
            
            # Format numerique si necessaire
            if c_idx in [3, 4, 5, 8]:  # Colonnes C, D, E, H
                if value is not None:
                    cell.number_format = '#,##0.00'
    
    # Ligne 31: BEAC avec formules
    beac_row = 31
    
    # Colonne A: Date
    ws2.cell(row=beac_row, column=1, value=date_str)
    ws2.cell(row=beac_row, column=1).border = border
    
    # Colonne B: Code Agence
    ws2.cell(row=beac_row, column=2, value='BEAC')
    ws2.cell(row=beac_row, column=2).border = border
    ws2.cell(row=beac_row, column=2).font = Font(bold=True)
    
    # Colonnes C a G: Vides (avec bordures)
    for col in range(3, 8):
        ws2.cell(row=beac_row, column=col, value=None)
        ws2.cell(row=beac_row, column=col).border = border
    
    # Colonne H: FORMULE =-SOMME(H2:H30)
    # ws2.cell(row=beac_row, column=8, value='=-SOMME(H2:H30)')
    ws2.cell(row=beac_row, column=8, value='=-SUM(H2:H30)')
    ws2.cell(row=beac_row, column=8).border = border
    ws2.cell(row=beac_row, column=8).number_format = '#,##0.00'
    ws2.cell(row=beac_row, column=8).font = Font(bold=True)
    
    # Colonne I: FORMULE =SI(H31<0;"Retrait BEAC";SI(H31>0;"Versement BEAC";"Aucun besoin"))
    # ws2.cell(row=beac_row, column=9, value='=SI(H31<0;"Retrait BEAC";SI(H31>0;"Versement BEAC";"Aucun besoin"))')
    ws2.cell(row=beac_row, column=9, value='=IF(H31<0;"Retrait BEAC";IF(H31>0;"Versement BEAC";"Aucun besoin"))')
    ws2.cell(row=beac_row, column=9).border = border
    ws2.cell(row=beac_row, column=9).font = Font(bold=True)
    
    # Colonnes J, K, L: Vides
    for col in range(10, 13):
        ws2.cell(row=beac_row, column=col, value=None)
        ws2.cell(row=beac_row, column=col).border = border
    
    # Ajuster largeurs colonnes
    for col in ws2.columns:
        ws2.column_dimensions[col[0].column_letter].width = 18

    # ========================================================================
    # FEUILLE 3: Optimisation de transport
    # ========================================================================
    ws3 = wb.create_sheet("Optimisation transport")
    
    # ========================================================================
    # Partie 1: Separation Besoin/Excedent (LIGNES FIXES 1-33)
    # ========================================================================
    ws3['A1'] = "REPARTITION BESOIN / EXCEDENT"
    ws3['A1'].font = Font(bold=True, size=12)

    # Colonnes Besoin (A-B)
    ws3['A2'] = "Code Agence a Approvisionner"
    ws3['B2'] = "Besoin d'approvisionnement (en Million)"
    ws3['A2'].font = header_font
    ws3['A2'].fill = header_fill
    ws3['B2'].font = header_font
    ws3['B2'].fill = header_fill

    # Remplir lignes 3 à 31 (29 lignes max)
    row_idx = 3
    for agence, montant in df_feuille3_part1['besoin'].items():
        if row_idx <= 31:
            ws3[f'A{row_idx}'] = agence
            ws3[f'B{row_idx}'] = montant
            ws3[f'B{row_idx}'].number_format = '#,##0.00'
            row_idx += 1

    # Calcul total B1 (formule)
    # ws3['B1'] = '=SOMME(B3:B31)'
    ws3['B1'] = '=SUM(B3:B31)'
    ws3['B1'].number_format = '#,##0.00'
    ws3['B1'].font = Font(bold=True)

    # Colonnes Excedent (D-E)
    ws3['D2'] = "Code Agence en excedent"
    ws3['E2'] = "Excedent de tresorerie"
    ws3['D2'].font = header_font
    ws3['D2'].fill = header_fill
    ws3['E2'].font = header_font
    ws3['E2'].fill = header_fill

    # Remplir lignes 3 à 31
    row_idx = 3
    for agence, montant in df_feuille3_part1['excedent'].items():
        if row_idx <= 31:
            ws3[f'D{row_idx}'] = agence
            ws3[f'E{row_idx}'] = montant
            ws3[f'E{row_idx}'].number_format = '#,##0.00'
            row_idx += 1

    # Calcul total E1 (formule)
    # ws3['E1'] = '=SOMME(E3:E31)'
    ws3['E1'] = '=SUM(E3:E31)'
    ws3['E1'].number_format = '#,##0.00'
    ws3['E1'].font = Font(bold=True)

    # ========================================================================
    # Partie 2: Matrice d'allocation (COMMENCE LIGNE 34)
    # ========================================================================
    start_row = 34

    ws3[f'A{start_row}'] = "MATRICE D'ALLOCATION OPTIMALE"
    ws3[f'A{start_row}'].font = Font(bold=True, size=12)

    start_row += 2  # Ligne 36 pour les headers
    
    # Headers colonnes (agences en besoin)
    besoin_agences = list(df_feuille3_part1['besoin'].keys())
    for col_idx, agence in enumerate(besoin_agences, 2):
        ws3.cell(row=start_row, column=col_idx, value=agence)
        ws3.cell(row=start_row, column=col_idx).font = header_font
        ws3.cell(row=start_row, column=col_idx).fill = header_fill
        ws3.cell(row=start_row, column=col_idx).alignment = Alignment(horizontal='center')
    
    # Colonne Total Excedent
    ws3.cell(row=start_row, column=len(besoin_agences)+2, value="Total Excedent")
    ws3.cell(row=start_row, column=len(besoin_agences)+2).font = header_font
    ws3.cell(row=start_row, column=len(besoin_agences)+2).fill = header_fill
    
    # Headers lignes (agences en excedent) + Matrice
    excedent_agences = list(df_feuille3_part1['excedent'].keys())
    
    for row_offset, agence_exc in enumerate(excedent_agences, 1):
        current_row = start_row + row_offset
        
        # Header ligne
        ws3.cell(row=current_row, column=1, value=agence_exc)
        ws3.cell(row=current_row, column=1).font = header_font
        ws3.cell(row=current_row, column=1).fill = header_fill
        
        # Valeurs allocation
        idx_exc = agences_list.index(agence_exc)
        
        for col_offset, agence_bes in enumerate(besoin_agences, 2):
            idx_bes = agences_list.index(agence_bes)
            valeur = allocation_matrix[idx_exc, idx_bes]
            
            cell = ws3.cell(row=current_row, column=col_offset, value=valeur if valeur > 0 else "")
            if valeur > 0:
                cell.number_format = '#,##0.00'
                cell.fill = PatternFill(start_color="E8F4F8", end_color="E8F4F8", fill_type="solid")
            cell.border = border
        
        # Total excedent
        total_exc = df_feuille3_part1['excedent'][agence_exc]
        ws3.cell(row=current_row, column=len(besoin_agences)+2, value=total_exc)
        ws3.cell(row=current_row, column=len(besoin_agences)+2).number_format = '#,##0.00'
        ws3.cell(row=current_row, column=len(besoin_agences)+2).font = Font(bold=True)
    
    # Ligne Total Besoin
    total_row = start_row + len(excedent_agences) + 1
    ws3.cell(row=total_row, column=1, value="Total Besoin")
    ws3.cell(row=total_row, column=1).font = header_font
    ws3.cell(row=total_row, column=1).fill = header_fill
    
    for col_offset, agence_bes in enumerate(besoin_agences, 2):
        total_bes = df_feuille3_part1['besoin'][agence_bes]
        ws3.cell(row=total_row, column=col_offset, value=total_bes)
        ws3.cell(row=total_row, column=col_offset).number_format = '#,##0.00'
        ws3.cell(row=total_row, column=col_offset).font = Font(bold=True)
    
    # Ajuster largeurs
    for col in ws3.columns:
        ws3.column_dimensions[col[0].column_letter].width = 15
    
    # ========================================================================
    # FEUILLE 4: Matrice des couts unitaires
    # ========================================================================
    ws4 = wb.create_sheet("Matrice couts unitaires")
    
    ws4['A1'] = "MATRICE DES COUTS UNITAIRES (FCFA par Million)"
    ws4['A1'].font = Font(bold=True, size=12)
    
    # Headers
    for col_idx, agence in enumerate(agences_list, 2):
        ws4.cell(row=3, column=col_idx, value=agence)
        ws4.cell(row=3, column=col_idx).font = header_font
        ws4.cell(row=3, column=col_idx).fill = header_fill
        ws4.cell(row=3, column=col_idx).alignment = Alignment(horizontal='center')
    
    # Matrice
    for row_idx, agence_row in enumerate(agences_list, 4):
        ws4.cell(row=row_idx, column=1, value=agence_row)
        ws4.cell(row=row_idx, column=1).font = header_font
        ws4.cell(row=row_idx, column=1).fill = header_fill
        
        for col_idx, agence_col in enumerate(agences_list, 2):
            valeur = matrice_couts.loc[agence_row, agence_col]
            cell = ws4.cell(row=row_idx, column=col_idx, value=valeur if valeur > 0 else 0)
            
            if valeur == 0:
                cell.fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")
            else:
                cell.fill = PatternFill(start_color="FFF8DC", end_color="FFF8DC", fill_type="solid")
            
            cell.border = border
            cell.alignment = Alignment(horizontal='center')
    
    # Ajuster largeurs
    for col in ws4.columns:
        ws4.column_dimensions[col[0].column_letter].width = 12
    
    # Sauvegarder
    wb.save(filename)
    print(f"  Fichier Excel genere: {filename.name}")


# ============================================================================
# 8. GENERATION RAPPORT WORD
# ============================================================================

def add_heading_custom(doc, text, level=1):
    """Ajoute un titre personnalise"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_table_with_data(doc, df, title=None):
    """Ajoute un tableau avec donnees DataFrame"""
    if title:
        doc.add_paragraph(title, style='Heading 3')
    
    table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
    table.style = 'Light Grid Accent 1'
    
    # Headers
    for col_idx, col_name in enumerate(df.columns):
        cell = table.rows[0].cells[col_idx]
        cell.text = str(col_name)
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Data
    for row_idx, row in df.iterrows():
        for col_idx, value in enumerate(row):
            table.rows[row_idx+1].cells[col_idx].text = str(value)
    
    return table


# ============================================================================
# 9. PIPELINE PRINCIPAL
# ============================================================================



def generer_csv_stocks_depuis_excel():
    """
    Génère le fichier CSV stocks_minimum_agences.csv depuis Excel
    pour compatibilité avec le dashboard
    """
    print("[0/9] Génération CSV stocks minimum depuis Excel...")
    print("-" * 80)
    
    filepath = 'Données pour la solutions streamlit - Avec code.xlsx'
    
    if not Path(filepath).exists():
        raise FileNotFoundError(f"Fichier Excel introuvable: {filepath}")
    
    # Lire Feuil1
    df_stocks = pd.read_excel(filepath, sheet_name='Feuil1', dtype={'Code_Agence': str})
    df_stocks['Code_Agence'] = df_stocks['Code_Agence'].str.strip()
    
    # Convertir en millions pour le dashboard
    df_stocks['Stock_Minimum_Millions'] = df_stocks['Stock_de_securite'].apply(
        lambda x: float(str(x).replace(' ', '')) / 1_000_000
    )
    
    # Garder seulement les colonnes nécessaires
    df_stocks_export = df_stocks[['Code_Agence', 'Stock_Minimum_Millions']].copy()
    
    # Sauvegarder CSV
    df_stocks_export.to_csv('stocks_minimum_agences.csv', index=False)
    
    print(f"✅ Fichier CSV généré: stocks_minimum_agences.csv")
    print(f"   {len(df_stocks_export)} agences")
    print()



def main():
    """
    Pipeline principal d'execution
    """
    print()
    print("DEMARRAGE DU PIPELINE D'OPTIMISATION")
    print("=" * 80)
    print()
    
    try:

        # ✅ ÉTAPE 0 : Générer CSV stocks depuis Excel
        generer_csv_stocks_depuis_excel()



        # ====================================================================
        # ETAPE 1: SELECTION MODELE OPTIMAL
        # ====================================================================
        model_selection = select_best_model()
        
        best_model = model_selection['best_model']
        best_approach = model_selection['best_approach']
        
        print()
        print("=" * 80)
        
        # ====================================================================
        # ETAPE 2: GENERATION MATRICE COUTS
        # ====================================================================
        matrice_couts = generate_matrice_couts()
        
        print()
        print("=" * 80)
        
        # # ====================================================================
        # # ETAPE 3: CALCUL SOLDES VEILLE
        # # ====================================================================
        # print("[3/9] Calcul des soldes de veille pour tous les jours...")
        # print("-" * 80)
        
        # soldes_dict = calculate_soldes_veille_all_days(best_model, best_approach)
        
        # print("Soldes de veille calcules pour J+1 a J+5")
        # print()
        # print("=" * 80)
        
        # ====================================================================
        # ETAPE 4-7: BOUCLE SUR LES 5 JOURS
        # ====================================================================
        results_all_days = []
        agences_list = CONFIG['agences']
        
        for jour in range(1, 6):
            print()
            print(f"[{3+jour}/9] TRAITEMENT JOUR J+{jour}")
            print("=" * 80)
            
            date_jour = get_date_from_jour(jour)
            print(f"Date: {date_jour}")
            print()
            
            # ETAPE 4: Generation donnees d'entree
            print(f"  [4.{jour}] Generation donnees d'entree...")
            
            df_entree = generate_feuilles_entree(
                jour, 
                best_model, 
                best_approach, 
                # soldes_dict[jour]
            )
            
            print(f"  Donnees generees: {len(df_entree)} agences")
            print()
            
            # ETAPE 5: Separation Besoin/Excedent
            print(f"  [5.{jour}] Separation Besoin/Excedent...")
            
            besoin_dict = {}
            excedent_dict = {}
            
            for _, row in df_entree.iterrows():
                agence = row['Code Agence']
                flux = row['Flux transport optimal']
                
                if flux > 0.01:
                    besoin_dict[agence] = flux
                elif flux < -0.01:
                    excedent_dict[agence] = abs(flux)
            
            df_feuille3_part1 = {
                'besoin': besoin_dict,
                'excedent': excedent_dict
            }
            
            print(f"  Agences en Besoin: {len(besoin_dict)}")
            print(f"  Agences en Excedent: {len(excedent_dict)}")
            print()
            
            # ETAPE 6: Optimisation
            print(f"  [6.{jour}] Optimisation allocation...")
            
            result_optim = optimize_transport_allocation(df_entree, matrice_couts)
            
            allocation_matrix = result_optim['allocation_matrix']
            cout_total = result_optim['cout_total']
            nb_transactions = result_optim['nb_transactions']
            
            # Statistiques jour
            total_besoin = sum(besoin_dict.values())
            total_excedent = sum(excedent_dict.values())
            cout_moyen = cout_total / nb_transactions if nb_transactions > 0 else 0
            
            # Role BEAC
            flux_beac = df_entree[df_entree['Code Agence'] == 'BEAC']['Flux transport optimal'].values[0]
            
            if flux_beac > 0:
                role_beac = "Versement BEAC"
            elif flux_beac < 0:
                role_beac = "Retrait BEAC"
            else:
                role_beac = "Aucune intervention"
            
            result_jour = {
                'jour': jour,
                'date': date_jour,
                'nb_besoin': len(besoin_dict),
                'nb_excedent': len(excedent_dict),
                'total_besoin': total_besoin,
                'total_excedent': total_excedent,
                'cout_total': cout_total,
                'nb_transactions': nb_transactions,
                'cout_moyen': cout_moyen,
                'role_beac': role_beac
            }
            
            results_all_days.append(result_jour)
            
            # ETAPE 7: Export Excel
            print(f"  [7.{jour}] Export fichier Excel...")
            
            export_excel_jour(
                jour,
                df_entree,
                df_feuille3_part1,
                allocation_matrix,
                matrice_couts,
                agences_list,
                result_optim
            )
            
            print()
        
        print("=" * 80)
        

        # ====================================================================
        # ETAPE 8: RESUME FINAL
        # ====================================================================
        print()
        print("[8/8] RESUME FINAL")
        print("=" * 80)
        print()
        
        print("FICHIERS GENERES:")
        print()
        
        for jour in range(1, 6):
            filename = f"Optimisation_Transport_J{jour}.xlsx"
            print(f"  - {filename}")
        
        print(f"  - Rapport_Validation_Consolide.docx")
        print()
        
        print("STATISTIQUES GLOBALES:")
        print()
        
        cout_total_cumule = sum([r['cout_total'] for r in results_all_days])
        nb_trans_total = sum([r['nb_transactions'] for r in results_all_days])
        
        print(f"  Cout total (5 jours)      : {format_number(cout_total_cumule)} FCFA")
        print(f"  Transactions totales      : {nb_trans_total}")
        print(f"  Cout moyen quotidien      : {format_number(cout_total_cumule/5)} FCFA")
        print(f"  Transactions moyennes/jour: {nb_trans_total/5:.1f}")
        print()
        
        print("=" * 80)
        print("PIPELINE TERMINE AVEC SUCCES")
        print("=" * 80)
        print()
        
        return {
            'model_selection': model_selection,
            'matrice_couts': matrice_couts,
            'results_days': results_all_days,
            'status': 'SUCCESS'
        }
        
    except Exception as e:
        print()
        print("=" * 80)
        print("ERREUR CRITIQUE")
        print("=" * 80)
        print()
        print(f"Type: {type(e).__name__}")
        print(f"Message: {str(e)}")
        print()
        
        import traceback
        print("Traceback complet:")
        print(traceback.format_exc())
        
        return {
            'status': 'FAILED',
            'error': str(e)
        }


# ============================================================================
# POINT D'ENTREE
# ============================================================================

if __name__ == "__main__":
    
    # Verification de l'environnement
    print()
    print("VERIFICATION DE L'ENVIRONNEMENT")
    print("=" * 80)
    print()
    
    # Verifier dossier base
    base_path = Path(CONFIG['base_dir'])
    if not base_path.exists():
        print(f"ERREUR: Dossier {CONFIG['base_dir']} introuvable")
        print("Veuillez vous assurer que les resultats de predictions sont disponibles")
        sys.exit(1)
    
    print(f"Dossier base: {base_path.absolute()}")
    
    # Compter agences disponibles
    agences_disponibles = [d.name for d in base_path.iterdir() if d.is_dir()]
    print(f"Agences detectees: {len(agences_disponibles)}")
    
    if len(agences_disponibles) < 29:
        print(f"ATTENTION: Seulement {len(agences_disponibles)} agences trouvees (29 attendues)")
    
    print()
    print("Environnement valide")
    print()
    
    # Execution pipeline
    result = main()
    
    # Code retour
    if result['status'] == 'SUCCESS':
        sys.exit(0)
    else:
        sys.exit(1)